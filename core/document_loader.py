"""
Document loader supporting multiple file formats:
  - Plain text (.txt)
  - PDF (.pdf) via pypdf
  - Word documents (.docx) via python-docx
  - Visio diagrams (.vsdx) via XML extraction from the OPC package
"""

import io
import zipfile
import xml.etree.ElementTree as ET


SUPPORTED_EXTENSIONS = ["txt", "pdf", "docx", "vsdx"]


def load_text(file_bytes: bytes, filename: str) -> str:
    """Route to the correct parser based on file extension."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    parsers = {
        "txt": _parse_txt,
        "pdf": _parse_pdf,
        "docx": _parse_docx,
        "vsdx": _parse_vsdx,
    }
    parser = parsers.get(ext)
    if parser is None:
        raise ValueError(f"Unsupported file type '.{ext}'. Supported: {SUPPORTED_EXTENSIONS}")
    return parser(file_bytes)


def _parse_txt(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    return data.decode("utf-8", errors="replace")


def _parse_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ImportError("Install pypdf to read PDFs:  pip install pypdf")

    reader = PdfReader(io.BytesIO(data))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        if text.strip():
            pages.append(f"[Page {i + 1}]\n{text.strip()}")
    if not pages:
        raise ValueError("Could not extract any text from the PDF. It may be image-based.")
    return "\n\n".join(pages)


def _parse_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Install python-docx to read Word files:  pip install python-docx")

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style and para.style.name.startswith("Heading"):
            level = para.style.name.replace("Heading", "").strip()
            prefix = "#" * (int(level) if level.isdigit() else 1)
            parts.append(f"{prefix} {text}")
        else:
            parts.append(text)

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    if not parts:
        raise ValueError("Could not extract any text from the Word document.")
    return "\n\n".join(parts)


def _parse_vsdx(data: bytes) -> str:
    """
    VSDX files are OPC (ZIP) packages containing XML.
    Text lives in visio/pages/page*.xml inside <vt:Text> and <Text> elements.
    We also check the document master shapes for label text.
    """
    texts: list[str] = []

    try:
        with zipfile.ZipFile(io.BytesIO(data), "r") as zf:
            xml_files = [
                n for n in zf.namelist()
                if n.endswith(".xml") and ("page" in n.lower() or "master" in n.lower())
            ]
            if not xml_files:
                xml_files = [n for n in zf.namelist() if n.endswith(".xml")]

            for xml_path in xml_files:
                try:
                    tree = ET.parse(zf.open(xml_path))
                    root = tree.getroot()
                    for elem in root.iter():
                        tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
                        if tag == "Text" and elem.text and elem.text.strip():
                            texts.append(elem.text.strip())
                        for child in elem:
                            if child.tail and child.tail.strip():
                                texts.append(child.tail.strip())
                except ET.ParseError:
                    continue
    except zipfile.BadZipFile:
        raise ValueError("Invalid VSDX file (not a valid ZIP/OPC package).")

    unique = list(dict.fromkeys(texts))
    if not unique:
        raise ValueError("Could not extract any text from the Visio file. "
                         "It may contain only shapes without text labels.")
    return "\n\n".join(unique)
